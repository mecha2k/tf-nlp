from dooly import Dooly
import docx
import warnings
import os

warnings.filterwarnings("ignore")
mt = Dooly(task="translation", lang="multi")

en_path = "../data/Finding Alphas.docx"
base = os.path.splitext(en_path)
ko_path = base[0] + "_ko" + base[1]

en_doc = docx.Document(en_path)
ko_doc = docx.Document()
print(len(en_doc.paragraphs))


for idx, paras in enumerate(en_doc.paragraphs):
    sentence = paras.text.strip()
    if idx == 1000:
        break
    if not sentence:
        continue
    result = mt(sentence, src_langs="en", tgt_langs="ko")
    para = ko_doc.add_paragraph(result)
    print(result)

ko_doc.save(ko_path)
