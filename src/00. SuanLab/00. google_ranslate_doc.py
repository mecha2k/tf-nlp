import docx
import os
import warnings
from googletrans import Translator


warnings.filterwarnings("ignore")

translator = Translator()
en_path = "../data/Finding Alphas.docx"
base = os.path.splitext(en_path)
ko_path = base[0] + "_ko" + base[1]

en_doc = docx.Document(en_path)
ko_doc = docx.Document()
print(len(en_doc.paragraphs))


for idx, paras in enumerate(en_doc.paragraphs):
    sentence = paras.text.strip()
    if idx == 1000:
        break
    if not sentence:
        continue
    result = translator.translate(sentence, src="en", dest="ko")
    para = ko_doc.add_paragraph(result.text)
    print(result.text)

ko_doc.save(ko_path)
